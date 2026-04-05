"""
RAG tizimi uchun dokumentlarni qayta ishlash va saqlash.
Fayl: backend/app/rag/processor.py
"""

import os
import json
import re
import hashlib
from pathlib import Path
from typing import List, Dict, Optional
import numpy as np

# ─── PAKETLAR ────────────────────────────────────────────
# pip install sentence-transformers PyPDF2 python-docx anthropic chromadb

try:
    from sentence_transformers import SentenceTransformer
    EMBED_AVAILABLE = True
except ImportError:
    EMBED_AVAILABLE = False
    print("Warning: sentence-transformers not installed. Run: pip install sentence-transformers")

try:
    import chromadb
    CHROMA_AVAILABLE = True
except ImportError:
    CHROMA_AVAILABLE = False
    print("Warning: chromadb not installed. Run: pip install chromadb")

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ─── MATN AJRATUVCHI ─────────────────────────────────────
class TextExtractor:
    """Har xil fayl turlaridan matn olish"""

    @staticmethod
    def from_txt(path: str) -> str:
        for enc in ["utf-8", "cp1251", "latin-1"]:
            try:
                with open(path, encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        return ""

    @staticmethod
    def from_pdf(path: str) -> str:
        if not PDF_AVAILABLE:
            return ""
        text = []
        try:
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        text.append(t)
        except Exception as e:
            print(f"PDF o'qish xatosi {path}: {e}")
        return "\n".join(text)

    @staticmethod
    def from_docx(path: str) -> str:
        if not DOCX_AVAILABLE:
            return ""
        try:
            doc = DocxDocument(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Jadvallardan ham matn ol
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            return "\n".join(paragraphs)
        except Exception as e:
            print(f"DOCX o'qish xatosi {path}: {e}")
            return ""

    @classmethod
    def extract(cls, path: str) -> str:
        ext = Path(path).suffix.lower()
        if ext == ".txt":
            return cls.from_txt(path)
        elif ext == ".pdf":
            return cls.from_pdf(path)
        elif ext in [".docx", ".doc"]:
            return cls.from_docx(path)
        return ""


# ─── MATN BO'LAKLARGA BO'LISH ────────────────────────────
class TextChunker:
    """Matnni kichik bo'laklarga bo'lish (overlap bilan)"""

    def __init__(self, chunk_size: int = 500, overlap: int = 50):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def split(self, text: str, source: str) -> List[Dict]:
        # Avval paragraflar bo'yicha ajrat
        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
        chunks = []
        current = []
        current_len = 0

        for para in paragraphs:
            words = para.split()
            if current_len + len(words) > self.chunk_size and current:
                chunk_text = " ".join(current)
                chunks.append({
                    "text": chunk_text,
                    "source": source,
                    "id": hashlib.md5(chunk_text.encode()).hexdigest()[:12]
                })
                # Overlap uchun oxirgi qismni saqla
                keep = current[-self.overlap:]
                current = keep
                current_len = len(keep)

            current.extend(words)
            current_len += len(words)

        if current:
            chunk_text = " ".join(current)
            chunks.append({
                "text": chunk_text,
                "source": source,
                "id": hashlib.md5(chunk_text.encode()).hexdigest()[:12]
            })

        return chunks


# ─── VECTOR STORE ─────────────────────────────────────────
class VectorStore:
    """ChromaDB asosida vector saqlash va qidirish"""

    def __init__(self, db_path: str = "./chroma_db"):
        if not CHROMA_AVAILABLE:
            raise ImportError("chromadb o'rnatilmagan")
        self.client = chromadb.PersistentClient(path=db_path)
        self.collection = self.client.get_or_create_collection(
            name="umidnoma_docs",
            metadata={"hnsw:space": "cosine"}
        )

    def add(self, chunks: List[Dict], embeddings: List[List[float]]):
        if not chunks:
            return
        self.collection.add(
            documents=[c["text"] for c in chunks],
            embeddings=embeddings,
            metadatas=[{"source": c["source"]} for c in chunks],
            ids=[c["id"] for c in chunks]
        )
        print(f"  ✓ {len(chunks)} ta chunk saqlandi")

    def search(self, query_embedding: List[float], n_results: int = 5) -> List[Dict]:
        results = self.collection.query(
            query_embeddings=[query_embedding],
            n_results=n_results
        )
        docs = []
        for i, doc in enumerate(results["documents"][0]):
            docs.append({
                "text": doc,
                "source": results["metadatas"][0][i]["source"],
                "distance": results["distances"][0][i]
            })
        return docs

    def count(self) -> int:
        return self.collection.count()


# ─── EMBEDDING MODEL ──────────────────────────────────────
class EmbeddingModel:
    """Multilingual sentence embedding"""

    def __init__(self, model_name: str = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"):
        if not EMBED_AVAILABLE:
            raise ImportError("sentence-transformers o'rnatilmagan")
        print(f"Embedding modeli yuklanmoqda: {model_name}")
        self.model = SentenceTransformer(model_name)
        print("✓ Model tayyor")

    def encode(self, texts: List[str]) -> List[List[float]]:
        embeddings = self.model.encode(texts, show_progress_bar=True)
        return embeddings.tolist()

    def encode_one(self, text: str) -> List[float]:
        return self.model.encode([text])[0].tolist()


# ─── RAG PIPELINE ─────────────────────────────────────────
class RAGProcessor:
    """To'liq RAG pipeline"""

    def __init__(self, docs_folder: str, db_path: str = "./chroma_db"):
        self.docs_folder = docs_folder
        self.chunker = TextChunker(chunk_size=500, overlap=50)
        self.embedder = EmbeddingModel()
        self.store = VectorStore(db_path=db_path)

    def process_all(self):
        """Barcha dokumentlarni o'qib, bazaga yoz"""
        docs_path = Path(self.docs_folder)
        if not docs_path.exists():
            print(f"Papka topilmadi: {self.docs_folder}")
            return

        files = list(docs_path.glob("**/*"))
        supported = [f for f in files if f.suffix.lower() in [".txt", ".pdf", ".docx", ".doc"]]

        print(f"\n{len(supported)} ta fayl topildi\n")

        all_chunks = []
        for file in supported:
            print(f"O'qilmoqda: {file.name}")
            text = TextExtractor.extract(str(file))
            if not text.strip():
                print(f"  ⚠ Matn topilmadi")
                continue
            chunks = self.chunker.split(text, source=file.name)
            print(f"  → {len(chunks)} ta chunk")
            all_chunks.extend(chunks)

        if not all_chunks:
            print("Hech qanday matn topilmadi!")
            return

        print(f"\nJami {len(all_chunks)} ta chunk embedding qilinmoqda...")
        texts = [c["text"] for c in all_chunks]
        embeddings = self.embedder.encode(texts)

        print("\nVector bazaga saqlanmoqda...")
        # 100 tadan batch qilib saqla
        batch_size = 100
        for i in range(0, len(all_chunks), batch_size):
            batch_chunks = all_chunks[i:i+batch_size]
            batch_embeddings = embeddings[i:i+batch_size]
            self.store.add(batch_chunks, batch_embeddings)

        print(f"\n✅ Tayyor! Jami {self.store.count()} ta chunk bazada")

    def query(self, question: str, n_results: int = 5) -> List[Dict]:
        """Savolga mos kontekst topish"""
        q_embedding = self.embedder.encode_one(question)
        results = self.store.search(q_embedding, n_results=n_results)
        return results


# ─── ISHGA TUSHIRISH ──────────────────────────────────────
if __name__ == "__main__":
    import sys

    # python processor.py /path/to/docs
    docs_folder = sys.argv[1] if len(sys.argv) > 1 else "./documents"

    processor = RAGProcessor(docs_folder=docs_folder)
    processor.process_all()
    print("\nTest so'rovi:")
    results = processor.query("stress bilan qanday kurashish kerak?")
    for r in results:
        print(f"\n[{r['source']}] (distance: {r['distance']:.3f})")
        print(r["text"][:200] + "...")
